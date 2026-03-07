"""Build PDF and DOCX bytes from QuoteRequest and QuoteAnalysis."""

from __future__ import annotations

import io
import tempfile
from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Image
from reportlab.platypus import Paragraph
from reportlab.platypus import SimpleDocTemplate
from reportlab.platypus import Spacer
from reportlab.platypus import Table
from reportlab.platypus import TableStyle

from models import QuoteAnalysis
from models import QuoteRequest

# Logo path: dockie_logo.svg in project root (parent of export/)
_LOGO_SVG_PATH = Path(__file__).resolve().parent.parent / "dockie_logo.svg"


def _logo_png_bytes() -> bytes | None:
    """Dockie logo as PNG bytes for embedding. Uses dockie_logo.png if present, else converts SVG (requires cairo)."""
    png_path = _LOGO_SVG_PATH.with_suffix(".png")
    if png_path.is_file():
        return png_path.read_bytes()
    if not _LOGO_SVG_PATH.is_file():
        return None
    try:
        import cairosvg
        png_io = io.BytesIO()
        cairosvg.svg2png(url=str(_LOGO_SVG_PATH), write_to=png_io)
        return png_io.getvalue()
    except Exception:
        return None


def _quote_request_from_state(raw: Any) -> QuoteRequest | None:
    if raw is None:
        return None
    if isinstance(raw, QuoteRequest):
        return raw
    if isinstance(raw, dict):
        return QuoteRequest.model_validate(raw)
    return None


def _quote_analysis_from_state(raw: Any) -> QuoteAnalysis | None:
    if raw is None:
        return None
    if isinstance(raw, QuoteAnalysis):
        return raw
    if isinstance(raw, dict):
        return QuoteAnalysis.model_validate(raw)
    return None


def build_pdf_bytes(
    quote_request: QuoteRequest | dict[str, Any] | None,
    quote_analysis: QuoteAnalysis | dict[str, Any] | None,
) -> bytes:
    """Build a one-page PDF quote document. Accepts Pydantic models or state dicts."""
    req = _quote_request_from_state(quote_request) if quote_request else None
    analysis = _quote_analysis_from_state(quote_analysis) if quote_analysis else None

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        rightMargin=inch,
        leftMargin=inch,
        topMargin=inch,
        bottomMargin=inch,
    )
    styles = getSampleStyleSheet()
    story = []
    logo_path = None  # temp file for logo PNG, cleaned up after build

    # Header: Dockie logo (banner); requires cairo for SVG→PNG, or add dockie_logo.png alongside the SVG
    logo_bytes = _logo_png_bytes()
    if logo_bytes:
        try:
            with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as f:
                f.write(logo_bytes)
                logo_path = f.name
            img = Image(logo_path, width=1.5 * inch, height=0.47 * inch)
            story.append(img)
            story.append(Spacer(1, 0.15 * inch))
        except Exception:
            if logo_path:
                Path(logo_path).unlink(missing_ok=True)
            logo_path = None

    title = Paragraph("Dockie Quote", styles["Title"])
    story.append(title)
    story.append(Spacer(1, 0.25 * inch))

    if req:
        story.append(Paragraph("Route & delivery details", styles["Heading2"]))
        data = [
            ["Origin", req.origin],
            ["Destination", req.destination],
            ["Vehicle type", req.vehicle_type],
            ["Delivery type", req.delivery_type],
        ]
        if req.item_to_ship:
            data.append(["Item(s) to ship", req.item_to_ship])
        if req.notes:
            data.append(["Notes", req.notes])
        t = Table(data, colWidths=[1.5 * inch, 4 * inch])
        t.setStyle(
            TableStyle(
                [
                    ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
                    ("FONTSIZE", (0, 0), (-1, -1), 10),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                ]
            )
        )
        story.append(t)
        story.append(Spacer(1, 0.2 * inch))

    if analysis:
        story.append(Paragraph("Price estimate", styles["Heading2"]))
        story.append(
            Paragraph(
                f"Matched scenario: {analysis.matched_scenario}",
                styles["Normal"],
            )
        )
        story.append(
            Paragraph(
                f"Estimated price: ${analysis.estimated_price:,.2f}",
                styles["Normal"],
            )
        )
        if analysis.breakdown:
            story.append(Paragraph(analysis.breakdown, styles["Normal"]))
        if analysis.confidence:
            story.append(
                Paragraph(f"Confidence: {analysis.confidence}", styles["Normal"])
            )

    doc.build(story)
    if logo_path:
        Path(logo_path).unlink(missing_ok=True)
    return buffer.getvalue()


def build_docx_bytes(
    quote_request: QuoteRequest | dict[str, Any] | None,
    quote_analysis: QuoteAnalysis | dict[str, Any] | None,
) -> bytes:
    """Build a DOCX quote document. Accepts Pydantic models or state dicts."""
    req = _quote_request_from_state(quote_request) if quote_request else None
    analysis = _quote_analysis_from_state(quote_analysis) if quote_analysis else None

    document = Document()
    # Header: Dockie logo (banner)
    logo_bytes = _logo_png_bytes()
    if logo_bytes:
        try:
            document.add_picture(io.BytesIO(logo_bytes), width=Inches(1.5))
        except Exception:
            pass
    document.add_heading("Dockie Quote", 0)

    if req:
        document.add_heading("Route & delivery details", level=1)
        document.add_paragraph(f"Origin: {req.origin}")
        document.add_paragraph(f"Destination: {req.destination}")
        document.add_paragraph(f"Vehicle type: {req.vehicle_type}")
        document.add_paragraph(f"Delivery type: {req.delivery_type}")
        if req.item_to_ship:
            document.add_paragraph(f"Item(s) to ship: {req.item_to_ship}")
        if req.notes:
            document.add_paragraph(f"Notes: {req.notes}")

    if analysis:
        document.add_heading("Price estimate", level=1)
        document.add_paragraph(f"Matched scenario: {analysis.matched_scenario}")
        p = document.add_paragraph()
        p.add_run("Estimated price: ").bold = True
        p.add_run(f"${analysis.estimated_price:,.2f}")
        if analysis.breakdown:
            document.add_paragraph(analysis.breakdown)
        if analysis.confidence:
            document.add_paragraph(f"Confidence: {analysis.confidence}")

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()
